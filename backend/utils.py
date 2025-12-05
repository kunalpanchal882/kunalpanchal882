"""
Utility functions for file handling and text extraction from resume files.
Supports PDF and DOCX formats using pdfplumber and python-docx libraries.
"""

import os
import pdfplumber
from docx import Document


def validate_file_type(filename: str) -> str | None:
    """
    Validates the file type based on file extension.
    
    Args:
        filename: The name of the uploaded file
    
    Returns:
        str: The file type ('pdf' or 'docx') if valid, None if unsupported
    """
    # Extract file extension
    if not filename:
        return None
    
    file_extension = filename.split('.')[-1].lower()
    
    # Return the file type if it's supported
    if file_extension in ['pdf', 'docx']:
        return file_extension
    
    return None


def extract_text_from_file(file_path: str, file_type: str) -> str:
    """
    Extracts text from a PDF or DOCX file.
    
    Args:
        file_path: The path to the file to extract text from
        file_type: The type of file ('pdf' or 'docx')
    
    Returns:
        str: The extracted text content from the file
    
    Raises:
        ValueError: If file type is unsupported or file cannot be read
        FileNotFoundError: If the file does not exist
    """
    
    # Check if file exists
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    # Extract text based on file type
    if file_type == "pdf":
        return _extract_text_from_pdf(file_path)
    
    elif file_type == "docx":
        return _extract_text_from_docx(file_path)
    
    else:
        raise ValueError(
            f"Unsupported file type: {file_type}. "
            "Please use PDF or DOCX files."
        )


def _extract_text_from_pdf(file_path: str) -> str:
    """
    Extracts text from a PDF file using pdfplumber.
    
    Args:
        file_path: Path to the PDF file
    
    Returns:
        str: Extracted text from all pages
    
    Raises:
        Exception: If PDF cannot be read
    """
    try:
        text_content = []
        
        # Open PDF and extract text from each page
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
        
        # Join all page text with newlines
        return "\n".join(text_content)
    
    except Exception as e:
        raise Exception(
            f"Failed to extract text from PDF file: {str(e)}"
        )


def _extract_text_from_docx(file_path: str) -> str:
    """
    Extracts text from a DOCX file using python-docx.
    
    Args:
        file_path: Path to the DOCX file
    
    Returns:
        str: Extracted text from all paragraphs
    
    Raises:
        Exception: If DOCX cannot be read
    """
    try:
        text_content = []
        
        # Open DOCX document
        document = Document(file_path)
        
        # Extract text from all paragraphs
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Also extract text from tables if present
        for table in document.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text)
                if row_text:
                    text_content.append(" ".join(row_text))
        
        # Join all text with newlines
        return "\n".join(text_content)
    
    except Exception as e:
        raise Exception(
            f"Failed to extract text from DOCX file: {str(e)}"
        )


def delete_file(file_path: str) -> bool:
    """
    Safely deletes a file if it exists.
    
    Args:
        file_path: Path to the file to delete
    
    Returns:
        bool: True if file was deleted, False if file did not exist
    """
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
            return True
        return False
    
    except Exception as e:
        # Log the error but don't raise - this is a cleanup operation
        print(f"Warning: Failed to delete temporary file {file_path}: {str(e)}")
        return False
